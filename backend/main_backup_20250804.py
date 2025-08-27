#!/usr/bin/env python3
"""
IntraNest 2.0 - Complete Document Management Edition - ENHANCED VERSION
FastAPI backend with Weaviate vector database, LlamaIndex RAG engine, and enhanced document processing
100% Operational - Production Ready with REAL-TIME PROGRESS TRACKING and ENHANCED PDF PROCESSING
"""

from fastapi import FastAPI, Depends, HTTPException, Header, File, UploadFile, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import StreamingResponse
from contextlib import asynccontextmanager
import os
import logging
from dotenv import load_dotenv
from typing import Optional, Dict, List, Any
import json
import uuid
import asyncio
from datetime import datetime, timedelta
from pathlib import Path
import hashlib
import tempfile
from io import BytesIO
import re
import traceback

# Enhanced PDF processing imports
import pdfplumber
import io
from typing import Dict, Any, Optional
import asyncio
from datetime import datetime

# Load environment variables first
load_dotenv()

# Configure logging with more detailed format
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('/tmp/intranest.log') if os.path.exists('/tmp') else logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# LlamaIndex imports - test each component individually
LLAMAINDEX_AVAILABLE = True
llamaindex_components = {}

try:
    from llama_index.core import VectorStoreIndex, Document, Settings, StorageContext
    from llama_index.core.node_parser import SentenceSplitter
    from llama_index.vector_stores.weaviate import WeaviateVectorStore
    from llama_index.llms.openai import OpenAI
    from llama_index.embeddings.openai import OpenAIEmbedding
    llamaindex_components.update({
        'VectorStoreIndex': VectorStoreIndex,
        'Document': Document,
        'Settings': Settings,
        'StorageContext': StorageContext,
        'SentenceSplitter': SentenceSplitter,
        'WeaviateVectorStore': WeaviateVectorStore,
        'OpenAI': OpenAI,
        'OpenAIEmbedding': OpenAIEmbedding
    })
    logger.info("✅ Core LlamaIndex components loaded")
except ImportError as e:
    logger.error(f"❌ Core LlamaIndex components failed: {e}")
    LLAMAINDEX_AVAILABLE = False

# Document processing imports
try:
    import weaviate
    from weaviate import WeaviateAsyncClient
    from weaviate.classes.config import Configure, Property, DataType
    from weaviate.classes.query import Filter
    import openai
    import redis
    import mimetypes
    # Enhanced PDF processing imports
    import docx
    from bs4 import BeautifulSoup
    import PyPDF2
    DOCUMENT_PROCESSING_AVAILABLE = True
    logger.info("✅ Document processing dependencies loaded (including enhanced PDF)")
except ImportError as e:
    logger.warning(f"⚠️ Document processing dependencies not available: {e}")
    DOCUMENT_PROCESSING_AVAILABLE = False

# Environment variables
WEAVIATE_URL = os.getenv("WEAVIATE_URL", "http://localhost:8080")
WEAVIATE_API_KEY = os.getenv("WEAVIATE_API_KEY", "b8f2c4e8a9d3f7e1b5c9a6e2f8d4b7c3e9a5f1d8b2c6e9f3a7b4e8c2d6f9a3b5c8")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
INTRANEST_API_KEY = os.getenv("INTRANEST_API_KEY")

logger.info(f"🔑 Environment check - OpenAI: {'✅' if OPENAI_API_KEY else '❌'}, Weaviate: {'✅' if WEAVIATE_API_KEY else '❌'}")

# Global service instances
document_service = None
storage_service = None
cache_service = None

# === ENHANCED DOCUMENT PROCESSING CLASSES ===

class EnhancedDocumentProcessor:
    """Enhanced document processor with PDF improvements and progress tracking"""
    
    def __init__(self):
        self.progress_cache = {}
    
    async def extract_text_from_pdf(self, file_content: bytes, filename: str, document_id: str) -> str:
        """Enhanced PDF text extraction using pdfplumber with progress tracking"""
        try:
            text_content = []
            
            # Use pdfplumber for better text extraction
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"📄 Processing PDF: {filename} ({total_pages} pages)")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Update progress (30% to 60% for text extraction)
                    progress = 30 + (page_num / total_pages) * 30
                    await self.update_progress(document_id, progress, f"Extracting text from page {page_num}/{total_pages}")
                    
                    # Extract text with better formatting
                    page_text = page.extract_text()
                    if page_text:
                        # Clean up text
                        cleaned_text = self.clean_extracted_text(page_text)
                        if cleaned_text.strip():
                            text_content.append(f"--- Page {page_num} ---\n{cleaned_text}\n")
                    
                    # Small delay to allow progress updates
                    await asyncio.sleep(0.01)
            
            final_text = "\n".join(text_content)
            
            if not final_text.strip():
                raise ValueError("No readable text found in PDF")
                
            logger.info(f"✅ PDF extraction successful: {len(final_text)} characters from {total_pages} pages")
            return final_text
            
        except Exception as e:
            logger.warning(f"⚠️ PDF extraction with pdfplumber failed for {filename}: {e}")
            # Fallback to PyPDF2 if pdfplumber fails
            return await self.fallback_pdf_extraction(file_content, filename, document_id)
    
    def clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts but preserve important characters
        text = re.sub(r'[^\w\s\.,!?;:()\-\'"@#$%^&*+=<>/\\|`~\[\]{}]', '', text)
        
        # Fix common OCR errors
        text = text.replace('ﬁ', 'fi').replace('ﬂ', 'fl')
        text = text.replace('–', '-').replace('—', '-')
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    async def fallback_pdf_extraction(self, file_content: bytes, filename: str, document_id: str) -> str:
        """Fallback PDF extraction using PyPDF2"""
        try:
            logger.info(f"🔄 Using PyPDF2 fallback for {filename}")
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = []
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                progress = 40 + (page_num / total_pages) * 20
                await self.update_progress(document_id, progress, f"Fallback extraction: page {page_num}/{total_pages}")
                
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    cleaned_text = self.clean_extracted_text(page_text)
                    if cleaned_text:
                        text_content.append(f"--- Page {page_num} ---\n{cleaned_text}\n")
            
            final_text = "\n".join(text_content) if text_content else "Unable to extract readable text from PDF"
            logger.info(f"✅ PyPDF2 fallback extraction: {len(final_text)} characters")
            return final_text
            
        except Exception as e:
            logger.error(f"❌ Fallback PDF extraction also failed for {filename}: {e}")
            return f"Error: Unable to extract text from PDF file '{filename}'. The file may be corrupted, encrypted, or contain only images."
    
    async def extract_text_from_docx(self, file_content: bytes, filename: str, document_id: str) -> str:
        """Extract text from DOCX files with progress tracking"""
        try:
            await self.update_progress(document_id, 30, "Extracting text from DOCX...")
            
            doc = docx.Document(io.BytesIO(file_content))
            text_content = []
            
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                
                if i % 10 == 0:  # Update progress every 10 paragraphs
                    progress = 30 + (i / total_paragraphs) * 30
                    await self.update_progress(document_id, progress, f"Processing paragraph {i+1}/{total_paragraphs}")
            
            final_text = "\n\n".join(text_content)
            logger.info(f"✅ DOCX extraction successful: {len(final_text)} characters from {total_paragraphs} paragraphs")
            return final_text
            
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed for {filename}: {e}")
            raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")
    
    async def extract_text_from_html(self, file_content: bytes, filename: str, document_id: str) -> str:
        """Extract text from HTML files with progress tracking"""
        try:
            await self.update_progress(document_id, 30, "Parsing HTML...")
            
            html_content = file_content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            final_text = '\n'.join(chunk for chunk in chunks if chunk)
            
            logger.info(f"✅ HTML extraction successful: {len(final_text)} characters")
            return final_text
            
        except Exception as e:
            logger.error(f"❌ HTML extraction failed for {filename}: {e}")
            raise ValueError(f"Failed to extract text from HTML file: {str(e)}")
    
    async def update_progress(self, document_id: str, progress: float, message: str):
        """Update processing progress"""
        if cache_service:
            cache_service.update_progress(
                document_id,
                "processing",
                min(100, max(0, progress)),
                message
            )
        
        self.progress_cache[document_id] = {
            'progress': min(100, max(0, progress)),
            'message': message,
            'timestamp': datetime.now().isoformat(),
            'status': 'processing' if progress < 100 else 'completed'
        }
        logger.info(f"📊 Progress {document_id}: {progress:.1f}% - {message}")

# === DOCUMENT MANAGEMENT CONFIGURATION ===
class DocumentConfig:
    """Document management configuration"""
    STORAGE_TYPE = os.getenv("STORAGE_TYPE", "local")
    UPLOAD_DIR = "/tmp/intranest_uploads"
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    ALLOWED_MIME_TYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain',
        'text/html',
        'text/markdown',
        'application/json'
    }
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200

class StorageService:
    """Handle file storage operations"""
    def __init__(self):
        try:
            os.makedirs(DocumentConfig.UPLOAD_DIR, exist_ok=True)
            logger.info(f"✅ Storage service initialized: {DocumentConfig.UPLOAD_DIR}")
        except Exception as e:
            logger.error(f"❌ Failed to initialize storage service: {e}")
            raise

    def save_uploaded_file(self, file_content: bytes, filename: str, user_id: str) -> str:
        """Save uploaded file locally and return file path"""
        try:
            # Create user-specific directory
            user_dir = os.path.join(DocumentConfig.UPLOAD_DIR, user_id)
            os.makedirs(user_dir, exist_ok=True)

            # Generate unique filename
            file_extension = Path(filename).suffix
            unique_filename = f"{uuid.uuid4().hex}{file_extension}"
            file_path = os.path.join(user_dir, unique_filename)

            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_content)

            logger.info(f"✅ File saved: {file_path} ({len(file_content)} bytes)")
            return file_path

        except Exception as e:
            logger.error(f"❌ Failed to save file: {e}")
            raise

class DocumentCacheService:
    """Handle document metadata caching with REAL progress tracking"""
    def __init__(self):
        self._memory_cache = {}
        self.document_metadata_cache = {}  # NEW: Cache for document listings
        self.use_redis = False

        try:
            redis_url = os.getenv("REDIS_URL", "redis://localhost:6379")
            self.redis_client = redis.Redis.from_url(redis_url, decode_responses=True)
            self.redis_client.ping()
            self.use_redis = True
            logger.info("✅ Redis cache connected")
        except Exception as e:
            logger.warning(f"⚠️ Redis not available, using memory cache: {e}")
            self.redis_client = None

    def cache_processing_status(self, document_id: str, status: Dict, ttl: int = 3600):
        """Cache document processing status"""
        try:
            if self.use_redis and self.redis_client:
                cache_key = f"processing:document:{document_id}"
                self.redis_client.setex(cache_key, ttl, json.dumps(status))
            else:
                self._memory_cache[document_id] = status
            logger.debug(f"📦 Cached status for {document_id}: {status.get('status', 'unknown')}")
        except Exception as e:
            logger.warning(f"⚠️ Cache write failed: {e}")
            self._memory_cache[document_id] = status

    def get_processing_status(self, document_id: str) -> Optional[Dict]:
        """Get document processing status"""
        try:
            if self.use_redis and self.redis_client:
                cache_key = f"processing:document:{document_id}"
                cached = self.redis_client.get(cache_key)
                return json.loads(cached) if cached else None
            else:
                return self._memory_cache.get(document_id)
        except Exception as e:
            logger.warning(f"⚠️ Cache read failed: {e}")
            return self._memory_cache.get(document_id)

    def update_progress(self, document_id: str, status: str, progress: int, message: str, **kwargs):
        """Update processing progress in real-time"""
        try:
            current_status = self.get_processing_status(document_id) or {}
            current_status.update({
                "status": status,
                "progress": progress,
                "message": message,
                "updated_at": datetime.now().isoformat(),
                **kwargs
            })
            self.cache_processing_status(document_id, current_status)
            logger.info(f"📊 Progress [{document_id}]: {progress}% - {message}")
        except Exception as e:
            logger.error(f"❌ Failed to update progress: {e}")

    def cache_document_metadata(self, user_id: str, document_id: str, metadata: Dict):
        """Cache document metadata for faster listings"""
        try:
            if user_id not in self.document_metadata_cache:
                self.document_metadata_cache[user_id] = {}
            
            # Enhanced metadata with proper data types
            enhanced_metadata = {
                "id": document_id,
                "filename": metadata.get("filename", "Unknown"),
                "size": int(metadata.get("file_size", 0)),
                "uploadDate": metadata.get("upload_date", datetime.now().isoformat()),
                "status": metadata.get("status", "completed"),
                "userId": user_id,
                "documentId": document_id,
                "chunks": int(metadata.get("chunks_created", 0)),
                "wordCount": int(metadata.get("word_count", 0)),
                "fileType": metadata.get("mime_type", "text/plain"),
                "file_size": int(metadata.get("file_size", 0)),  # Add both formats
                "upload_date": metadata.get("upload_date", datetime.now().isoformat()),
                "processing_status": metadata.get("status", "completed")
            }
            
            self.document_metadata_cache[user_id][document_id] = enhanced_metadata
            logger.info(f"📦 Cached metadata for {document_id}")
            
        except Exception as e:
            logger.warning(f"⚠️ Failed to cache document metadata: {e}")

    def get_user_documents(self, user_id: str) -> List[Dict]:
        """Get cached documents for user"""
        try:
            user_docs = self.document_metadata_cache.get(user_id, {})
            return list(user_docs.values())
        except Exception as e:
            logger.warning(f"⚠️ Failed to get cached documents: {e}")
            return []

class WeaviateHelper:
    """Helper class for Weaviate operations with robust property access"""

    @staticmethod
    def get_client():
        """Get Weaviate client with proper connection"""
        try:
            client = weaviate.connect_to_local(
                host="localhost",
                port=8080,
                auth_credentials=weaviate.auth.AuthApiKey(WEAVIATE_API_KEY) if WEAVIATE_API_KEY else None
            )
            logger.debug("✅ Weaviate client created")
            return client
        except Exception as e:
            logger.error(f"❌ Failed to connect to Weaviate: {e}")
            raise

    @staticmethod
    def safe_get_property(props, prop_name: str, default=''):
        """Safely get property value using multiple access methods"""
        try:
            # Method 1: Dictionary access
            if isinstance(props, dict):
                return props.get(prop_name, default)

            # Method 2: Attribute access
            if hasattr(props, prop_name):
                value = getattr(props, prop_name, default)
                return value if value is not None else default

            # Method 3: Check if it has __dict__
            if hasattr(props, '__dict__'):
                return props.__dict__.get(prop_name, default)

            return default

        except Exception as e:
            logger.warning(f"⚠️ Error accessing property {prop_name}: {e}")
            return default

    @staticmethod
    def safe_query_with_filter(collection, user_id: str, limit: int = 1000):
        """Query with user filter - handles different Weaviate versions"""
        try:
            logger.debug(f"🔍 Querying Weaviate for user: {user_id}")
            response = collection.query.fetch_objects(limit=limit)

            # Manual filtering with robust property access
            filtered_objects = []
            for obj in response.objects:
                try:
                    props = obj.properties if hasattr(obj, 'properties') else {}
                    obj_user_id = WeaviateHelper.safe_get_property(props, 'user_id', '')

                    if obj_user_id == user_id:
                        filtered_objects.append(obj)
                except Exception as e:
                    logger.warning(f"⚠️ Error filtering object: {e}")
                    continue

            response.objects = filtered_objects
            logger.debug(f"✅ Filtered to {len(filtered_objects)} objects for user: {user_id}")
            return response

        except Exception as e:
            logger.error(f"❌ Weaviate query failed: {e}")
            raise

class LlamaIndexRAGService:
    """RAG service with working hybrid search"""

    def __init__(self):
        self.weaviate_client = None
        self.vector_store = None
        self.index = None
        self.llm = None
        self.embed_model = None
        self.initialize_services()

    def initialize_services(self):
        """Initialize LlamaIndex services"""
        try:
            logger.info("🚀 Initializing LlamaIndex RAG service...")

            if not LLAMAINDEX_AVAILABLE or not llamaindex_components:
                logger.error("❌ LlamaIndex components not available")
                return

            if not OPENAI_API_KEY:
                logger.error("❌ OPENAI_API_KEY not found")
                return

            # Initialize Weaviate client
            try:
                self.weaviate_client = WeaviateHelper.get_client()
                if not self.weaviate_client.is_ready():
                    logger.error("❌ Weaviate client not ready")
                    return
                logger.info("✅ Weaviate client connected")
            except Exception as e:
                logger.error(f"❌ Weaviate connection failed: {e}")
                return

            # Initialize LlamaIndex components
            OpenAI = llamaindex_components['OpenAI']
            OpenAIEmbedding = llamaindex_components['OpenAIEmbedding']
            WeaviateVectorStore = llamaindex_components['WeaviateVectorStore']
            VectorStoreIndex = llamaindex_components['VectorStoreIndex']
            StorageContext = llamaindex_components['StorageContext']

            # Create LLM and embedding model
            self.llm = OpenAI(
                model="gpt-4",
                api_key=OPENAI_API_KEY,
                temperature=0.1,
                max_tokens=1500
            )

            self.embed_model = OpenAIEmbedding(
                model="text-embedding-3-small",
                api_key=OPENAI_API_KEY
            )

            # Configure Settings
            from llama_index.core import Settings
            Settings.llm = self.llm
            Settings.embed_model = self.embed_model

            # Setup Weaviate schema
            self.setup_weaviate_schema()

            # Connect to existing Weaviate vector store
            self.vector_store = WeaviateVectorStore(
                weaviate_client=self.weaviate_client,
                index_name="Documents",
                text_key="content"
            )

            # Create storage context
            storage_context = StorageContext.from_defaults(vector_store=self.vector_store)

            # Load or create index
            try:
                self.index = VectorStoreIndex.from_vector_store(
                    self.vector_store,
                    storage_context=storage_context
                )
                logger.info("✅ Connected to existing LlamaIndex")
            except Exception as e:
                logger.warning(f"⚠️ Creating new index: {e}")
                self.index = VectorStoreIndex(
                    nodes=[],
                    storage_context=storage_context
                )

            logger.info("✅ LlamaIndex RAG service initialized successfully")

        except Exception as e:
            logger.error(f"❌ Failed to initialize LlamaIndex RAG service: {e}")
            self.weaviate_client = None

    def setup_weaviate_schema(self):
        """Setup Weaviate schema"""
        try:
            collections = self.weaviate_client.collections.list_all()
            collection_names = [col for col in collections.keys()] if hasattr(collections, 'keys') else [col.name for col in collections]

            if "Documents" not in collection_names:
                logger.info("📋 Creating Documents collection...")

                from weaviate.classes.config import Configure

                collection = self.weaviate_client.collections.create(
                    name="Documents",
                    vectorizer_config=Configure.Vectorizer.text2vec_openai(
                        model="text-embedding-3-small"
                    ),
                    generative_config=Configure.Generative.openai(
                        model="gpt-4"
                    ),
                    properties=[
                        Property(name="content", data_type=DataType.TEXT),
                        Property(name="filename", data_type=DataType.TEXT),
                        Property(name="user_id", data_type=DataType.TEXT),
                        Property(name="document_id", data_type=DataType.TEXT),
                        Property(name="node_id", data_type=DataType.TEXT),
                        Property(name="chunk_id", data_type=DataType.INT),
                        Property(name="page_number", data_type=DataType.INT),
                        Property(name="metadata", data_type=DataType.OBJECT)
                    ]
                )
                logger.info("✅ Documents collection created")
            else:
                logger.info("✅ Documents collection already exists")

        except Exception as e:
            logger.error(f"❌ Schema setup error: {e}")

    async def search_documents(self, query: str, user_id: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search documents with hybrid approach and robust property access"""
        try:
            logger.info(f"🔍 Searching documents for query: '{query[:50]}...' user: {user_id}")

            if not self.weaviate_client:
                logger.error("❌ No Weaviate client available")
                return []

            documents_collection = self.weaviate_client.collections.get("Documents")
            response = documents_collection.query.near_text(
                query=query,
                limit=limit * 3  # Get more results to filter by user
            )

            logger.debug(f"📄 Weaviate returned {len(response.objects)} total results")

            filtered_results = []
            for obj in response.objects:
                try:
                    props = obj.properties if hasattr(obj, 'properties') else {}
                    obj_user_id = WeaviateHelper.safe_get_property(props, 'user_id', '')

                    if obj_user_id == user_id:
                        # Calculate similarity score
                        similarity_score = 0.8  # Default
                        if hasattr(obj, 'metadata') and obj.metadata:
                            certainty = getattr(obj.metadata, 'certainty', None)
                            distance = getattr(obj.metadata, 'distance', None)

                            if certainty is not None:
                                similarity_score = float(certainty)
                            elif distance is not None:
                                similarity_score = max(0.0, 1.0 - float(distance))

                        result = {
                            "content": WeaviateHelper.safe_get_property(props, 'content', ''),
                            "filename": WeaviateHelper.safe_get_property(props, 'filename', 'Unknown'),
                            "chunk_id": WeaviateHelper.safe_get_property(props, 'chunk_id', 0),
                            "page_number": WeaviateHelper.safe_get_property(props, 'page_number', 1),
                            "similarity_score": float(similarity_score),
                            "document_id": WeaviateHelper.safe_get_property(props, 'document_id', ''),
                            "node_id": str(obj.uuid)
                        }

                        filtered_results.append(result)
                        logger.debug(f"✅ Added result: {result['filename']} (score: {similarity_score:.3f})")

                        if len(filtered_results) >= limit:
                            break

                except Exception as e:
                    logger.warning(f"⚠️ Error processing search result: {e}")
                    continue

            logger.info(f"✅ Search found {len(filtered_results)} documents for user {user_id}")
            return filtered_results

        except Exception as e:
            logger.error(f"❌ Document search error: {e}")
            return []

    async def generate_rag_response(self, query: str, user_id: str, context_limit: int = 5) -> Dict[str, Any]:
        """Generate RAG response with document context"""
        try:
            logger.info(f"🔍 RAG query: {query[:100]}...")

            search_results = await self.search_documents(query, user_id, context_limit)

            if not search_results:
                return {
                    "response": f"""I don't have any documents in your knowledge base that are relevant to: "{query}"

To get started, please upload some documents using the document upload feature. I can then provide detailed answers based on your specific content.""",
                    "sources": [],
                    "has_context": False
                }

            # Build context from search results
            context_parts = []
            sources = []

            for i, result in enumerate(search_results):
                context_parts.append(f"[Source {i+1}] {result['content']}")

                similarity_score = result.get("similarity_score", 0.8)
                try:
                    relevance = round(float(similarity_score), 3)
                except (ValueError, TypeError):
                    relevance = 0.800

                sources.append({
                    "filename": result["filename"],
                    "page": result["page_number"],
                    "chunk": result["chunk_id"],
                    "relevance": relevance,
                    "node_id": result["node_id"]
                })

            context = "\n\n".join(context_parts)

            # Generate response using LLM
            if self.llm:
                try:
                    from llama_index.core.llms import ChatMessage, MessageRole

                    messages = [
                        ChatMessage(
                            role=MessageRole.SYSTEM,
                            content="You are IntraNest AI, a professional enterprise knowledge assistant. Provide accurate, well-structured responses based on the provided document context."
                        ),
                        ChatMessage(
                            role=MessageRole.USER,
                            content=f"""Based on the following documents from the user's knowledge base, answer the question: "{query}"

Context from documents:
{context}

Instructions:
1. Provide a comprehensive answer based on the provided context
2. Reference specific sources when making claims
3. If the context doesn't fully answer the question, acknowledge what's missing
4. Use a professional, clear tone suitable for business use
5. Structure your response with headers and bullet points when appropriate

Answer:"""
                        )
                    ]

                    response = self.llm.chat(messages)
                    ai_response = str(response)
                    logger.info("✅ Generated LLM response")
                except Exception as e:
                    logger.warning(f"⚠️ LLM response failed, using context summary: {e}")
                    ai_response = f"Based on your documents, here's what I found:\n\n{context[:500]}..."
            else:
                ai_response = f"Based on your documents:\n\n{context[:500]}..."

            # Format final response with sources
            if sources:
                sources_text = "\n\n**Sources:**\n"
                for source in sources:
                    sources_text += f"• {source['filename']} (Page {source['page']}, Relevance: {source['relevance']})\n"
                final_response = ai_response + sources_text
            else:
                final_response = ai_response

            return {
                "response": final_response,
                "sources": sources,
                "has_context": True,
                "context_chunks": len(sources)
            }

        except Exception as e:
            logger.error(f"❌ RAG response error: {e}")
            return {
                "response": f"I encountered an error while processing your query: {str(e)}",
                "sources": [],
                "has_context": False
            }

class ProfessionalResponseGenerator:
    """Enterprise-grade response generator"""

    def analyze_user_intent(self, message: str) -> Dict:
        """Analyze user intent from message"""
        message_lower = message.lower().strip()

        greeting_words = ["hello", "hi", "hey", "good morning", "good afternoon", "good evening"]
        if any(word in message_lower for word in greeting_words) and len(message.split()) <= 3:
            return {"type": "greeting", "confidence": 0.95}

        help_patterns = ["help", "what can you do", "capabilities", "features", "how to use", "what do you do"]
        if any(phrase in message_lower for phrase in help_patterns):
            return {"type": "help", "confidence": 0.9}

        return {"type": "contextual", "confidence": 0.7}

    def get_time_appropriate_greeting(self) -> str:
        """Get time-appropriate greeting"""
        current_hour = datetime.now().hour
        if 5 <= current_hour < 12:
            return "Good morning"
        elif 12 <= current_hour < 17:
            return "Good afternoon"
        elif 17 <= current_hour < 22:
            return "Good evening"
        else:
            return "Hello"

    def generate_greeting_response(self) -> str:
        """Professional greeting with time awareness"""
        greeting = self.get_time_appropriate_greeting()
        return f"""{greeting}! I'm IntraNest AI, your enterprise knowledge assistant.

I can help you with document analysis, knowledge search, data insights, and content creation. All responses include source citations when referencing your organizational knowledge.

What would you like to work on today?"""

    def generate_help_response(self) -> str:
        """Professional capabilities overview"""
        return """**IntraNest AI Capabilities**

I can assist you with:

**📄 Document Analysis**
- Summarize reports, policies, and technical documents
- Extract key insights and action items
- Compare multiple documents using advanced RAG

**🔍 Knowledge Search**
- Find specific information across your knowledge base
- Answer questions with precise source citations
- Provide contextual explanations using semantic search

**📊 Data Insights**
- Analyze trends and patterns in your documents
- Generate executive summaries with source attribution
- Create actionable recommendations based on your content

What specific task would you like assistance with?"""

    async def generate_professional_response(self, user_message: str, user_id: str = "anonymous", model: str = "IntraNest-AI") -> str:
        """Generate enterprise-grade responses"""
        try:
            if not user_message or user_message.strip() == "":
                user_message = "Hello"

            intent = self.analyze_user_intent(user_message)
            logger.debug(f"🎯 Intent detected: {intent['type']}")

            # For contextual queries, try RAG first
            if intent["type"] == "contextual" and document_service:
                rag_result = await document_service.generate_rag_response(user_message, user_id)
                if rag_result["has_context"]:
                    logger.info("✅ Using RAG response with document context")
                    return rag_result["response"]

            # Fall back to standard responses
            if intent["type"] == "greeting":
                return self.generate_greeting_response()
            elif intent["type"] == "help":
                return self.generate_help_response()
            else:
                return f"I can help you with that topic. Please upload some documents so I can provide specific information based on your content."

        except Exception as e:
            logger.error(f"❌ Response generation error: {e}")
            return "I apologize, but I encountered an error processing your request. Please try again."

# Initialize response generator
response_generator = ProfessionalResponseGenerator()

# === ENHANCED DOCUMENT PROCESSING FUNCTIONS ===

def create_chunks_with_progress(text_content: str, filename: str, document_id: str, user_id: str, cache_service, progress_start: int = 40) -> List[Dict]:
    """Create text chunks with progress updates"""
    try:
        if not text_content.strip():
            logger.warning(f"⚠️ Empty text content for {document_id}")
            return []

        logger.info(f"🔪 Creating chunks for {document_id}: {len(text_content)} characters")

        chunks = []
        chunk_size = DocumentConfig.CHUNK_SIZE
        overlap = DocumentConfig.CHUNK_OVERLAP

        cache_service.update_progress(
            document_id,
            "chunking",
            progress_start,
            "Breaking document into chunks..."
        )

        # Create chunks
        for i in range(0, len(text_content), chunk_size - overlap):
            chunk_text = text_content[i:i + chunk_size].strip()
            if not chunk_text:
                continue

            chunk_data = {
                "content": chunk_text,
                "filename": filename,
                "user_id": user_id,
                "document_id": document_id,
                "node_id": str(uuid.uuid4()),
                "chunk_id": len(chunks),
                "page_number": 1,
                "metadata": {
                    "upload_date": datetime.now().isoformat(),
                    "file_type": mimetypes.guess_type(filename)[0] or "text/plain",
                    "chunk_size": len(chunk_text),
                    "source": filename
                }
            }
            chunks.append(chunk_data)

        logger.info(f"✅ Created {len(chunks)} chunks for {document_id}")

        cache_service.update_progress(
            document_id,
            "chunks_created",
            progress_start + 15,
            f"Created {len(chunks)} text chunks"
        )

        return chunks

    except Exception as e:
        logger.error(f"❌ Chunking failed for {document_id}: {e}")
        raise

async def store_chunks_with_progress(chunks: List[Dict], document_id: str, cache_service, progress_start: int = 55) -> int:
    """Store chunks in Weaviate with progress updates"""
    client = None
    try:
        logger.info(f"📊 Storing {len(chunks)} chunks for {document_id}")

        client = WeaviateHelper.get_client()
        documents_collection = client.collections.get("Documents")
        success_count = 0
        total_chunks = len(chunks)

        for i, chunk in enumerate(chunks):
            try:
                # Insert chunk
                result = documents_collection.data.insert(chunk)
                success_count += 1

                logger.debug(f"✅ Stored chunk {chunk['chunk_id']} -> {result}")

                # Update progress
                progress = progress_start + int(((i + 1) / total_chunks) * 35)  # 35% for storage
                cache_service.update_progress(
                    document_id,
                    "storing",
                    progress,
                    f"Stored {success_count}/{total_chunks} chunks",
                    chunks_processed=success_count,
                    total_chunks=total_chunks
                )

                await asyncio.sleep(0.01)  # Small delay for progress visibility

            except Exception as e:
                logger.warning(f"⚠️ Failed to store chunk {i}: {e}")
                continue

        logger.info(f"📊 Successfully stored {success_count}/{total_chunks} chunks")
        return success_count

    except Exception as e:
        logger.error(f"❌ Storage failed: {e}")
        raise
    finally:
        if client:
            try:
                client.close()
            except:
                pass

async def process_document_with_enhanced_extraction(document_id: str, document_metadata: Dict, file_content: bytes, cache_service):
    """Process document with enhanced text extraction and real-time progress tracking"""
    processor = EnhancedDocumentProcessor()
    
    try:
        filename = document_metadata["filename"]
        user_id = document_metadata["user_id"]

        logger.info(f"📄 Processing document: {filename} for user: {user_id}")

        # Step 1: Text extraction with enhanced processing (30% -> 60%)
        cache_service.update_progress(document_id, "extracting_text", 30, "Starting text extraction...")

        try:
            filename_lower = filename.lower()
            
            if filename_lower.endswith('.pdf'):
                text_content = await processor.extract_text_from_pdf(file_content, filename, document_id)
            elif filename_lower.endswith('.docx'):
                text_content = await processor.extract_text_from_docx(file_content, filename, document_id)
            elif filename_lower.endswith(('.txt', '.md')):
                await processor.update_progress(document_id, 30, "Reading text file...")
                text_content = file_content.decode('utf-8', errors='ignore')
            elif filename_lower.endswith(('.html', '.htm')):
                text_content = await processor.extract_text_from_html(file_content, filename, document_id)
            elif filename_lower.endswith('.json'):
                await processor.update_progress(document_id, 30, "Parsing JSON...")
                text_content = file_content.decode('utf-8', errors='ignore')
            else:
                # Default text extraction
                await processor.update_progress(document_id, 30, "Reading file as text...")
                text_content = file_content.decode('utf-8', errors='ignore')

            logger.info(f"📊 Extracted {len(text_content)} characters from {filename}")

        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            raise Exception(f"Failed to extract text from {filename}: {str(e)}")

        if not text_content.strip():
            raise Exception(f"No readable text found in {filename}")

        cache_service.update_progress(document_id, "text_extracted", 60, f"Extracted {len(text_content)} characters")

        # Calculate word count for metadata
        word_count = len(text_content.split())

        # Step 2: Create chunks (60% -> 75%)
        logger.info(f"🔪 Creating chunks for {document_id}")
        chunks = create_chunks_with_progress(text_content, filename, document_id, user_id, cache_service, 60)

        if not chunks:
            raise Exception("No content chunks created")

        # Step 3: Store chunks (75% -> 95%)
        logger.info(f"💾 Storing chunks for {document_id}")
        success_count = await store_chunks_with_progress(chunks, document_id, cache_service, 75)

        if success_count == 0:
            raise Exception("Failed to store any chunks")

        # Step 4: Finalize (95% -> 100%)
        cache_service.update_progress(document_id, "finalizing", 95, "Finalizing...")

        # Update metadata with enhanced information
        document_metadata.update({
            "chunks_created": success_count,
            "status": "completed",
            "total_chunks": len(chunks),
            "chunks_processed": success_count,
            "text_length": len(text_content),
            "word_count": word_count
        })

        # Cache document metadata for fast listings
        cache_service.cache_document_metadata(user_id, document_id, document_metadata)

        logger.info(f"🎉 Document processing completed: {success_count} chunks stored, {word_count} words")

    except Exception as e:
        logger.error(f"❌ Document processing failed: {e}")
        cache_service.update_progress(document_id, "error", 0, f"Processing failed: {str(e)}")
        document_metadata.update({"status": "error", "error": str(e)})
        raise

# === INITIALIZATION ===

async def initialize_services():
    """Initialize all services"""
    global storage_service, cache_service, document_service

    try:
        logger.info("🚀 Initializing IntraNest services...")

        # Initialize storage service
        storage_service = StorageService()
        logger.info("✅ Storage service initialized")

        # Initialize cache service
        cache_service = DocumentCacheService()
        logger.info("✅ Cache service initialized")

        # Initialize document service
        if DOCUMENT_PROCESSING_AVAILABLE and LLAMAINDEX_AVAILABLE:
            document_service = LlamaIndexRAGService()
            logger.info("✅ Document service initialized")
        else:
            logger.warning("⚠️ Document service not available - missing dependencies")

        logger.info("✅ All services initialized successfully")

    except Exception as e:
        logger.error(f"❌ Service initialization failed: {e}")
        raise

# API Key dependency
async def verify_api_key(authorization: Optional[str] = Header(None)):
    """Verify API key from Authorization header"""
    if not INTRANEST_API_KEY:
        return True

    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header required")

    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization format")

    provided_key = authorization.replace("Bearer ", "")
    if provided_key != INTRANEST_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")

    return True

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan management"""
    logger.info("🚀 Starting IntraNest 2.0...")
    await initialize_services()
    logger.info("✅ IntraNest 2.0 started successfully")

    yield

    logger.info("🛑 Shutting down IntraNest 2.0...")
    if document_service and document_service.weaviate_client:
        try:
            document_service.weaviate_client.close()
        except:
            pass
    logger.info("✅ Shutdown complete")

# Create FastAPI application
app = FastAPI(
    title="IntraNest 2.0 - Enterprise AI Knowledge Management",
    description="Complete document processing and RAG system with enhanced PDF processing and real-time progress tracking",
    version="2.0.0-enhanced",
    lifespan=lifespan
)

# Add middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3090", "http://localhost:3080", "*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"]
)
app.add_middleware(GZipMiddleware, minimum_size=1000)

# === API ENDPOINTS ===

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "name": "IntraNest 2.0 - Enterprise AI Knowledge Management",
        "version": "2.0.0-enhanced",
        "status": "operational",
        "features": {
            "document_upload": "enabled",
            "enhanced_pdf_processing": "enabled",
            "document_processing": "enabled",
            "document_listing": "enabled",
            "rag_search": "enabled",
            "real_time_progress": "enabled"
        },
        "services": {
            "storage": storage_service is not None,
            "cache": cache_service is not None,
            "document_processing": document_service is not None,
            "weaviate": document_service and document_service.weaviate_client is not None
        }
    }

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "services": {
            "api": "running",
            "storage": "operational" if storage_service else "unavailable",
            "cache": "operational" if cache_service else "unavailable",
            "document_service": "operational" if document_service else "unavailable",
            "weaviate": "operational" if document_service and document_service.weaviate_client else "unavailable"
        },
        "version": "2.0.0-enhanced"
    }

def generate_tokens(text: str):
    """Generator for streaming tokens"""
    sentences = text.split('. ')
    for i, sentence in enumerate(sentences):
        if i < len(sentences) - 1:
            sentence += '. '
        words = sentence.split(' ')
        for j, word in enumerate(words):
            if j == 0 and i == 0:
                yield word
            else:
                yield f" {word}"

async def generate_intranest_content(user_message: str, user_id: str = "anonymous", model: str = "IntraNest-AI") -> str:
    """Generate professional content"""
    return await response_generator.generate_professional_response(user_message, user_id, model)

@app.post("/chat/completions")
async def chat_completions_with_streaming(request: dict, api_key: bool = Depends(verify_api_key)):
    """OpenAI-compatible chat completions with streaming"""
    try:
        messages = request.get("messages", [])
        model = request.get("model", "IntraNest-AI")
        stream = request.get("stream", False)
        user_id = request.get("user_id", "anonymous")

        # FIX: Try to extract user_id from various sources
        if user_id == "anonymous":
            # Try to get from request headers or other sources
            # For now, we'll check if there are recent documents and use the most common user_id
            try:
                client = WeaviateHelper.get_client()
                try:
                    documents_collection = client.collections.get("Documents")
                    recent_response = documents_collection.query.fetch_objects(limit=50)

                    # Count user_ids in recent documents
                    user_counts = {}
                    for obj in recent_response.objects:
                        props = obj.properties if hasattr(obj, 'properties') else {}
                        if isinstance(props, dict):
                            obj_user_id = props.get('user_id', '')
                            if obj_user_id and obj_user_id != 'anonymous':
                                user_counts[obj_user_id] = user_counts.get(obj_user_id, 0) + 1

                    # Use the most common user_id (likely the current user)
                    if user_counts:
                        user_id = max(user_counts, key=user_counts.get)
                        logger.info(f"🔄 Auto-detected user_id: {user_id} from recent documents")

                finally:
                    client.close()
            except Exception as e:
                logger.warning(f"⚠️ Failed to auto-detect user_id: {e}")

        if not messages:
            messages = [{"role": "user", "content": "Hello"}]

        user_message = "Hello"
        for msg in reversed(messages):
            if msg.get("role") == "user" and msg.get("content"):
                user_message = msg.get("content")
                break

        logger.info(f"💬 Chat request: '{user_message[:100]}...' user: {user_id}")

        response_content = await generate_intranest_content(user_message, user_id, model)

        if stream:
            async def event_generator():
                try:
                    completion_id = f"chatcmpl-{int(datetime.now().timestamp())}-{uuid.uuid4().hex[:8]}"
                    created = int(datetime.now().timestamp())

                    for token in generate_tokens(response_content):
                        chunk = {
                            "id": completion_id,
                            "object": "chat.completion.chunk",
                            "created": created,
                            "model": model,
                            "choices": [{
                                "index": 0,
                                "delta": {"content": token},
                                "finish_reason": None
                            }]
                        }
                        yield f"data: {json.dumps(chunk)}\n\n"
                        await asyncio.sleep(0.01)

                    final_chunk = {
                        "id": completion_id,
                        "object": "chat.completion.chunk",
                        "created": created,
                        "model": model,
                        "choices": [{
                            "index": 0,
                            "delta": {},
                            "finish_reason": "stop"
                        }]
                    }
                    yield f"data: {json.dumps(final_chunk)}\n\n"
                    yield "data: [DONE]\n\n"

                except Exception as e:
                    logger.error(f"❌ Streaming error: {e}")
                    yield f"data: [DONE]\n\n"

            return StreamingResponse(
                event_generator(),
                media_type="text/plain",
                headers={
                    "Cache-Control": "no-cache",
                    "Connection": "keep-alive"
                }
            )
        else:
            return {
                "id": f"chatcmpl-{int(datetime.now().timestamp())}-{uuid.uuid4().hex[:8]}",
                "object": "chat.completion",
                "created": int(datetime.now().timestamp()),
                "model": model,
                "choices": [{
                    "index": 0,
                    "message": {
                        "role": "assistant",
                        "content": response_content
                    },
                    "finish_reason": "stop"
                }],
                "usage": {
                    "prompt_tokens": max(1, len(user_message) // 4),
                    "completion_tokens": max(1, len(response_content) // 4),
                    "total_tokens": max(2, len(user_message + response_content) // 4)
                }
            }

    except Exception as e:
        logger.error(f"❌ Chat completions error: {e}")
        return {
            "id": f"chatcmpl-error-{int(datetime.now().timestamp())}",
            "object": "chat.completion",
            "created": int(datetime.now().timestamp()),
            "model": "IntraNest-AI",
            "choices": [{
                "index": 0,
                "message": {
                    "role": "assistant",
                    "content": "I apologize, but I encountered an error. Please try again."
                },
                "finish_reason": "stop"
            }],
            "usage": {"prompt_tokens": 1, "completion_tokens": 20, "total_tokens": 21}
        }

# === ENHANCED DOCUMENT MANAGEMENT ENDPOINTS ===

@app.post("/api/documents/upload")
async def upload_document_with_enhanced_processing(
    file: UploadFile = File(...),
    user_id: str = Form(...),
    api_key: bool = Depends(verify_api_key)
):
    """Upload and process document with enhanced PDF processing and real-time progress tracking"""

    # Service availability check
    if not storage_service or not cache_service:
        logger.error("❌ Required services not available")
        raise HTTPException(status_code=503, detail="Document services not available")

    try:
        logger.info(f"📤 Enhanced upload started: {file.filename} for user: {user_id}")

        # File validation
        file_size = file.size or 0
        if file_size > DocumentConfig.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds {DocumentConfig.MAX_FILE_SIZE // (1024*1024)}MB limit"
            )

        mime_type = file.content_type or mimetypes.guess_type(file.filename)[0]
        if mime_type not in DocumentConfig.ALLOWED_MIME_TYPES:
            raise HTTPException(
                status_code=400,
                detail="File type not supported. Use PDF, DOCX, TXT, HTML, or MD files."
            )

        # Generate document ID
        document_id = f"doc_{int(datetime.now().timestamp())}_{uuid.uuid4().hex[:8]}"
        logger.info(f"📋 Document ID: {document_id}")

        # Step 1: Initialize (5%)
        cache_service.update_progress(
            document_id, "initializing", 5, "Starting enhanced upload...",
            filename=file.filename, user_id=user_id, file_size=file_size
        )

        # Step 2: Read file (15%)
        cache_service.update_progress(document_id, "reading", 15, "Reading file...")
        file_content = await file.read()
        logger.info(f"📖 Read {len(file_content)} bytes")

        # Step 3: Save file (25%)
        cache_service.update_progress(document_id, "saving", 25, "Saving file...")
        file_path = storage_service.save_uploaded_file(file_content, file.filename, user_id)

        # Create enhanced metadata
        document_metadata = {
            "document_id": document_id,
            "filename": file.filename,
            "user_id": user_id,
            "file_path": file_path,
            "status": "processing",
            "created_at": datetime.now().isoformat(),
            "file_size": len(file_content),
            "mime_type": mime_type,
            "processing_started_at": datetime.now().isoformat(),
            "upload_date": datetime.now().isoformat()
        }

        cache_service.cache_processing_status(document_id, document_metadata)

        # Step 4: Process document with enhanced extraction (25% -> 95%)
        logger.info(f"🚀 Starting enhanced processing for {document_id}")
        try:
            await process_document_with_enhanced_extraction(document_id, document_metadata, file_content, cache_service)

            # Step 5: Complete (100%)
            cache_service.update_progress(
                document_id, "completed", 100,
                f"Processing complete! Created {document_metadata.get('chunks_created', 0)} chunks.",
                chunks_created=document_metadata.get('chunks_created', 0)
            )

            document_metadata["status"] = "completed"
            document_metadata["completed_at"] = datetime.now().isoformat()
            cache_service.cache_processing_status(document_id, document_metadata, ttl=86400)

            logger.info(f"✅ Enhanced upload completed: {file.filename} -> {document_metadata.get('chunks_created', 0)} chunks")

            return {
                "success": True,
                "document_id": document_id,
                "filename": file.filename,
                "status": "completed",
                "message": f"Successfully processed '{file.filename}' with enhanced extraction",
                "chunks_created": document_metadata.get("chunks_created", 0),
                "word_count": document_metadata.get("word_count", 0)
            }

        except Exception as process_error:
            logger.error(f"❌ Enhanced processing failed: {process_error}")

            cache_service.update_progress(
                document_id, "error", 0, f"Processing failed: {str(process_error)}"
            )

            document_metadata.update({
                "status": "error",
                "error": str(process_error),
                "failed_at": datetime.now().isoformat()
            })
            cache_service.cache_processing_status(document_id, document_metadata)

            return {
                "success": False,
                "document_id": document_id,
                "filename": file.filename,
                "status": "error",
                "error": str(process_error)
            }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Enhanced upload failed: {e}")
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Enhanced upload failed: {str(e)}")

@app.get("/api/documents/status/{document_id}")
async def get_processing_status(document_id: str, user_id: str, api_key: bool = Depends(verify_api_key)):
    """Get document processing status"""
    if not cache_service:
        raise HTTPException(status_code=503, detail="Cache service not available")

    try:
        status = cache_service.get_processing_status(document_id)
        if not status:
            raise HTTPException(status_code=404, detail="Document not found")

        if status.get("user_id") != user_id:
            raise HTTPException(status_code=403, detail="Access denied")

        return {
            "success": True,
            "document_id": document_id,
            "status": status.get("status", "unknown"),
            "progress": status.get("progress", 0),
            "message": status.get("message", "Processing..."),
            "filename": status.get("filename"),
            "created_at": status.get("created_at"),
            "updated_at": status.get("updated_at"),
            "completed_at": status.get("completed_at"),
            "error": status.get("error"),
            "chunks_created": status.get("chunks_created", 0),
            "word_count": status.get("word_count", 0)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Status check error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/documents/list")
async def list_documents_enhanced(request: dict, api_key: bool = Depends(verify_api_key)):
    """ENHANCED: List user documents with improved data mapping and caching"""
    try:
        user_id = request.get("user_id")
        if not user_id:
            raise HTTPException(status_code=400, detail="user_id is required")

        logger.info(f"📋 Enhanced listing documents for user: {user_id}")

        # First, try to get from cache
        cached_documents = cache_service.get_user_documents(user_id) if cache_service else []
        
        if cached_documents:
            logger.info(f"✅ Returning {len(cached_documents)} cached documents")
            return {
                "success": True,
                "documents": cached_documents,
                "total": len(cached_documents),
                "userId": user_id,
                "source": "cache"
            }

        # If no cache, query Weaviate
        client = WeaviateHelper.get_client()
        try:
            documents_collection = client.collections.get("Documents")
            response = documents_collection.query.fetch_objects(limit=5000)

            logger.info(f"📄 Raw query returned {len(response.objects)} total objects")

            # Group by document_id with enhanced data mapping
            doc_map = {}
            processed_count = 0
            user_ids_found = set()

            for obj in response.objects:
                try:
                    props = obj.properties if hasattr(obj, 'properties') else {}

                    # Enhanced property access
                    if isinstance(props, dict):
                        obj_user_id = props.get('user_id', '')
                        user_ids_found.add(obj_user_id)

                        # Filter by user
                        if obj_user_id == user_id:
                            processed_count += 1

                            # Get document properties with safe defaults
                            doc_id = props.get('document_id', str(obj.uuid))
                            filename = props.get('filename', 'Unknown')
                            content = props.get('content', '')
                            metadata = props.get('metadata', {})

                            if doc_id not in doc_map:
                                # Handle metadata safely with enhanced mapping
                                metadata_dict = metadata if isinstance(metadata, dict) else {}
                                
                                # Enhanced document object with all required fields
                                doc_map[doc_id] = {
                                    "id": doc_id,
                                    "filename": filename,
                                    "size": metadata_dict.get('file_size', 0) or 0,
                                    "uploadDate": metadata_dict.get('upload_date', datetime.now().isoformat()),
                                    "status": "processed",
                                    "userId": user_id,
                                    "documentId": doc_id,
                                    "chunks": 0,
                                    "wordCount": 0,
                                    "fileType": metadata_dict.get('file_type', 'text/plain'),
                                    # Additional fields for better UI display
                                    "file_size": metadata_dict.get('file_size', 0) or 0,
                                    "upload_date": metadata_dict.get('upload_date', datetime.now().isoformat()),
                                    "processing_status": "completed"
                                }

                            # Accumulate stats
                            doc_map[doc_id]["chunks"] += 1
                            doc_map[doc_id]["size"] += len(content) if isinstance(content, str) else 0
                            doc_map[doc_id]["wordCount"] += len(content.split()) if isinstance(content, str) else 0

                except Exception as e:
                    logger.warning(f"⚠️ Error processing document object: {e}")
                    continue

            documents = list(doc_map.values())
            
            # Cache the results for future requests
            if cache_service:
                for doc in documents:
                    cache_service.cache_document_metadata(user_id, doc["id"], doc)

            logger.info(f"✅ Enhanced listing: {processed_count} chunks, {len(documents)} documents for user {user_id}")

            return {
                "success": True,
                "documents": documents,
                "total": len(documents),
                "userId": user_id,
                "source": "weaviate",
                "debug": {
                    "total_objects": len(response.objects),
                    "processed_chunks": processed_count,
                    "unique_documents": len(documents),
                    "user_ids_found": list(user_ids_found)
                }
            }

        finally:
            client.close()

    except Exception as e:
        logger.error(f"❌ Enhanced list documents error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")

@app.post("/api/documents/search")
async def search_documents_endpoint(request: dict, api_key: bool = Depends(verify_api_key)):
    """Search documents using RAG"""
    try:
        query = request.get("query", "")
        user_id = request.get("user_id", "")
        limit = request.get("limit", 5)

        if not query or not user_id:
            raise HTTPException(status_code=400, detail="query and user_id are required")

        if not document_service:
            raise HTTPException(status_code=503, detail="Document service not available")

        results = await document_service.search_documents(query, user_id, limit)

        return {
            "success": True,
            "query": query,
            "results": results,
            "total": len(results)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/chat/rag")
async def rag_chat_endpoint(request: dict, api_key: bool = Depends(verify_api_key)):
    """RAG-powered chat endpoint"""
    try:
        query = request.get("query", "")
        user_id = request.get("user_id", "anonymous")
        context_limit = request.get("context_limit", 5)

        if not query:
            raise HTTPException(status_code=400, detail="query is required")

        if not document_service:
            raise HTTPException(status_code=503, detail="Document service not available")

        result = await document_service.generate_rag_response(query, user_id, context_limit)

        return {
            "success": True,
            "query": query,
            "response": result["response"],
            "sources": result["sources"],
            "has_context": result["has_context"],
            "context_chunks": result.get("context_chunks", 0)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ RAG chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/documents/delete")
async def delete_document_endpoint(request: dict, api_key: bool = Depends(verify_api_key)):
    """Delete a document and all its chunks"""
    try:
        document_id = request.get("document_id", "")
        user_id = request.get("user_id", "")

        if not document_id or not user_id:
            raise HTTPException(status_code=400, detail="document_id and user_id are required")

        logger.info(f"🗑️ Deleting document: {document_id} for user: {user_id}")

        client = WeaviateHelper.get_client()
        try:
            documents_collection = client.collections.get("Documents")
            
            # First, verify the document belongs to the user
            response = documents_collection.query.fetch_objects(limit=1000)
            user_objects = []
            
            for obj in response.objects:
                props = obj.properties if hasattr(obj, 'properties') else {}
                obj_user_id = WeaviateHelper.safe_get_property(props, 'user_id', '')
                obj_doc_id = WeaviateHelper.safe_get_property(props, 'document_id', '')
                
                if obj_user_id == user_id and obj_doc_id == document_id:
                    user_objects.append(obj.uuid)

            if not user_objects:
                raise HTTPException(status_code=404, detail="Document not found or access denied")

            # Delete all chunks for this document
            deleted_count = 0
            for obj_uuid in user_objects:
                try:
                    documents_collection.data.delete_by_id(obj_uuid)
                    deleted_count += 1
                except Exception as e:
                    logger.warning(f"⚠️ Failed to delete chunk {obj_uuid}: {e}")

            # Remove from cache
            if cache_service and hasattr(cache_service, 'document_metadata_cache'):
                user_cache = cache_service.document_metadata_cache.get(user_id, {})
                if document_id in user_cache:
                    del user_cache[document_id]

            logger.info(f"✅ Deleted {deleted_count} chunks for document {document_id}")

            return {
                "success": True,
                "message": f"Deleted document and {deleted_count} chunks",
                "deleted_chunks": deleted_count
            }

        finally:
            client.close()

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Delete error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# === PROGRESS TRACKING ENDPOINT ===

@app.get("/api/documents/progress/{document_id}")
async def get_upload_progress(document_id: str, api_key: bool = Depends(verify_api_key)):
    """Get real-time upload progress"""
    if not cache_service:
        return {
            "progress": 0,
            "message": "Cache service not available",
            "status": "unknown",
            "timestamp": datetime.now().isoformat()
        }
    
    try:
        status = cache_service.get_processing_status(document_id)
        if status:
            return {
                "success": True,
                "progress": status.get("progress", 0),
                "status": status.get("status", "unknown"),
                "message": status.get("message", "Processing..."),
                "timestamp": status.get("updated_at", datetime.now().isoformat()),
                "chunks_created": status.get("chunks_created", 0),
                "chunks_processed": status.get("chunks_processed", 0),
                "total_chunks": status.get("total_chunks", 0)
            }
        else:
            return {
                "success": False,
                "progress": 0,
                "message": "Processing not started or completed",
                "status": "unknown",
                "timestamp": datetime.now().isoformat()
            }
    except Exception as e:
        logger.error(f"❌ Progress check error: {e}")
        return {
            "success": False,
            "progress": 0,
            "message": f"Error checking progress: {str(e)}",
            "status": "error",
            "timestamp": datetime.now().isoformat()
        }

# === DEBUG ENDPOINTS ===

@app.get("/api/debug/services")
async def debug_services(api_key: bool = Depends(verify_api_key)):
    """Debug service status"""
    return {
        "services": {
            "storage_service": storage_service is not None,
            "cache_service": cache_service is not None,
            "document_service": document_service is not None,
            "weaviate_client": document_service and document_service.weaviate_client is not None
        },
        "environment": {
            "openai_key_present": bool(OPENAI_API_KEY),
            "weaviate_key_present": bool(WEAVIATE_API_KEY),
            "intranest_key_present": bool(INTRANEST_API_KEY)
        },
        "capabilities": {
            "llamaindex_available": LLAMAINDEX_AVAILABLE,
            "document_processing_available": DOCUMENT_PROCESSING_AVAILABLE,
            "enhanced_pdf_processing": True,
            "pdfplumber_available": True,
            "docx_processing": True,
            "html_processing": True
        }
    }

@app.get("/api/debug/weaviate-inspect")
async def inspect_weaviate_data(api_key: bool = Depends(verify_api_key)):
    """Inspect raw Weaviate data for debugging"""
    try:
        client = WeaviateHelper.get_client()
        try:
            documents_collection = client.collections.get("Documents")
            response = documents_collection.query.fetch_objects(limit=5)

            objects_info = []
            for i, obj in enumerate(response.objects):
                obj_info = {
                    "index": i,
                    "uuid": str(obj.uuid),
                    "properties_type": str(type(obj.properties)),
                    "has_properties": hasattr(obj, 'properties')
                }

                # Try to extract info safely using our robust method
                try:
                    props = obj.properties
                    obj_info["user_id"] = WeaviateHelper.safe_get_property(props, 'user_id', 'NOT_FOUND')
                    obj_info["filename"] = WeaviateHelper.safe_get_property(props, 'filename', 'NOT_FOUND')
                    obj_info["document_id"] = WeaviateHelper.safe_get_property(props, 'document_id', 'NOT_FOUND')

                    # Get property structure info
                    if isinstance(props, dict):
                        obj_info["property_keys"] = list(props.keys())
                    else:
                        obj_info["property_attributes"] = [attr for attr in dir(props) if not attr.startswith('_')]

                except Exception as e:
                    obj_info["error"] = str(e)

                objects_info.append(obj_info)

            return {
                "success": True,
                "total_objects": len(response.objects),
                "objects": objects_info
            }

        finally:
            client.close()

    except Exception as e:
        logger.error(f"❌ Weaviate inspection failed: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/debug/test-enhanced-upload")
async def test_enhanced_upload(
    file: UploadFile = File(...),
    user_id: str = Form(...),
    api_key: bool = Depends(verify_api_key)
):
    """Test enhanced upload processing without storing"""
    try:
        logger.info(f"🧪 Test enhanced upload: {file.filename} for {user_id}")

        file_content = await file.read()
        logger.info(f"🧪 Read {len(file_content)} bytes")

        # Test enhanced text extraction
        processor = EnhancedDocumentProcessor()
        document_id = f"test_{uuid.uuid4().hex[:8]}"
        
        filename_lower = file.filename.lower()
        
        if filename_lower.endswith('.pdf'):
            text_content = await processor.extract_text_from_pdf(file_content, file.filename, document_id)
        elif filename_lower.endswith('.docx'):
            text_content = await processor.extract_text_from_docx(file_content, file.filename, document_id)
        elif filename_lower.endswith(('.html', '.htm')):
            text_content = await processor.extract_text_from_html(file_content, file.filename, document_id)
        else:
            text_content = file_content.decode('utf-8', errors='ignore')

        return {
            "success": True,
            "filename": file.filename,
            "size": len(file_content),
            "user_id": user_id,
            "text_length": len(text_content),
            "word_count": len(text_content.split()),
            "text_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
            "processing_method": "enhanced_extraction"
        }
    except Exception as e:
        logger.error(f"❌ Test enhanced upload failed: {e}")
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
